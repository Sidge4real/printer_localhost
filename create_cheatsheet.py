from docx import Document

# Document aanmaken
doc = Document()
doc.add_heading("Overzicht commando's en uitleg - Printen via WSL Canon MG5100", 0)

# Commando's + uitleg
commands_explanations = [
    ("usbipd list", "Toont alle aangesloten USB apparaten op Windows die je kunt delen met WSL."),
    ("usbipd bind --busid 2-3", "Bind het USB-apparaat met busid 2-3 zodat het gedeeld kan worden met WSL."),
    ("usbipd attach --busid 2-3 --wsl", "Verbindt het gebonden USB apparaat met WSL."),
    ("sudo service cups status", "Controleer of de CUPS printservice draait in WSL."),
    ("sudo service cups restart", "Herstart de CUPS service, bijvoorbeeld bij problemen."),
    ("sudo lpinfo -v", "Toont beschikbare printerpoorten en apparaten."),
    ("sudo lpinfo -m | grep -i mg5100", "Zoekt naar printerdrivers die bij de MG5100 passen."),
    ('sudo lpadmin -p MG5100 -E -v "usb://Canon/MG5100%20series?serial=306BCF&interface=1" -m gutenprint.5.3://bjc-PIXMA-MG5100/expert',
     "Voegt de printer MG5100 toe aan CUPS met de juiste driver."),
    ("lpstat -p MG5100", "Geeft status van de printer MG5100 weer."),
    ("lpstat -p MG5100 -l", "Geeft gedetailleerde status van de printer."),
    ("lp -d MG5100 test.txt", "Stuur het bestand test.txt naar printer MG5100 om te printen."),
    ("sudo tail -n 50 /var/log/cups/error_log", "Bekijk de laatste 50 regels van het CUPS logboek voor foutanalyse."),
    ("sudo tail -f /var/log/cups/error_log", "Bekijk realtime updates van het CUPS logboek."),
    ("sudo apt update", "Werk pakketlijsten bij in Ubuntu."),
    ("sudo apt install nodejs npm", "Installeer Node.js en npm in Ubuntu."),
    ("npm init -y", "Initialiseer een nieuw Node.js project."),
    ("npm install express multer", "Installeer benodigde npm pakketten voor de print backend."),
    ("node index.js", "Start de Node.js print backend server."),
    ("lp -d MG5100 /mnt/c/Users/vande/Documents/test.pdf", "Print een bestand vanaf Windows-schijf (via WSL pad)."),
    ("wsl lp -d MG5100 /mnt/c/Users/vande/Documents/test.pdf", "Print een Windows bestand via WSL vanuit PowerShell.")
]

# Commando's toevoegen aan document
for cmd, explanation in commands_explanations:
    doc.add_paragraph(cmd, style='ListBullet')
    doc.add_paragraph(explanation)
    doc.add_paragraph("")  # lege regel voor overzicht

# Opslaan in Windows-map
filepath = "printen_via_wsl_commando_en_uitleg.docx"
doc.save(filepath)

print(f"Document succesvol opgeslagen als: {filepath}")
